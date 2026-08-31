#!/usr/bin/env -S uv run
"""Real-world regression scenarios for skull, run against the LIVE Qwen
endpoint (and E2B if configured) - not mocked.

The pytest suite in tests/ mocks every network call, which is correct for
unit-testing individual functions but has repeatedly missed real bugs that
only show up when the actual model decides how to use the actual tools:
the stale-bytecode-after-overwrite bug, the forget() silent-paraphrase-
failure bug, the base64-through-the-conversation context blowup, the
embedding-similarity-threshold miscalibration, and the wholesale-skill-
exclusion-in-plan-mode gap were ALL found this way, not by unit tests.

Run every scenario:
    uv run tests/scenarios.py

Run one scenario:
    uv run tests/scenarios.py --only 2

Each scenario is a real multi-step task modeled on ones actually run
during development, chosen to exercise several subsystems together (the
way a real bug always did) rather than one function in isolation:

  1. Document + skill chain - read a real .docx, call several existing
     skills on facts extracted from it, verify the results are numerically
     correct. Catches: document-extraction bugs, wrong-tool-call bugs,
     skill dispatch bugs.
  2. Skill lifecycle - create a skill, run it, overwrite it with
     materially different logic, verify the NEW logic actually takes
     effect (not stale bytecode), roll back, verify the ORIGINAL logic is
     restored. Catches: the stale-bytecode-after-overwrite bug, the
     no-rollback-on-a-bad-edit gap.
  3. Memory contradiction + fuzzy forget - state a preference, restate a
     contradicting one in very different wording, verify the old one is
     superseded (not left duplicated), then ask to forget a different
     fact by paraphrase rather than exact quote. Catches: the embedding-
     similarity-threshold-too-high bug, forget()'s silent-failure-on-
     paraphrase bug.
  4. Plan-mode boundary - a multi-step task mixing read-only skills with
     one mutating skill and an explicit write request, all in plan mode;
     verify the write is blocked and a plan is produced instead of a
     false claim of completion, then /auto and verify the same action
     actually executes. Catches: the wholesale-skill-exclusion-in-plan-
     mode bug. NOTE: this one has been observed to fail intermittently -
     the model sometimes declines to call a genuinely available read-only
     skill in plan mode despite the system prompt saying it should, which
     is model sampling variance on this specific instruction rather than a
     deterministic bug in build_tools_and_impls (confirmed: build_tools_
     and_impls itself always classifies/includes the skill correctly -
     verified directly, independent of this script, whenever this failed).
     A single failure here is worth a re-run before treating it as a
     regression; a *consistent* failure across several runs is real.
  5. Long/heavy turn - a task that pulls a meaningful amount of real
     content through several tool calls in one turn, verifying no
     context-window 400 and that compaction (if it fires) leaves the
     conversation in a valid state. Catches: the base64-through-the-
     conversation context-overflow bug, the compaction-blind-spot bugs.
  6. Skill credential stays out of the model - create a skill that needs
     an API key, verify it's declared via required_env (not a plain
     parameter), verify request_skill_env is used to set it, and verify
     the actual secret value never appears anywhere in the conversation
     the model can see. Catches: the exact real incident this was built
     for - a model asking the user to paste an SMTP password directly
     into chat, landing it in plain text in conversation history/memory.

Isolation: SKILLS_DIR/PIPELINES_DIR/MEMORY_DIR are redirected to a fresh
temp directory before any skull module is imported, so this NEVER touches
your real skills/, pipelines/, or memory/ - only the live Qwen (and E2B,
if configured) endpoints are real. Files created in /tmp by a scenario are
removed at the end of that scenario, pass or fail.
"""

import argparse
import shutil
import sys
import tempfile
import traceback
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src"))

_TEMP_ROOT = Path(tempfile.mkdtemp(prefix="skull_scenarios_"))

import skull.config as _cfg  # noqa: E402

_cfg.SKILLS_DIR = _TEMP_ROOT / "skills"
_cfg.MEMORY_DIR = _TEMP_ROOT / "memory"
_cfg.PIPELINES_DIR = _TEMP_ROOT / "pipelines"
_cfg.SKILLS_DIR.mkdir(parents=True, exist_ok=True)
_cfg.MEMORY_DIR.mkdir(parents=True, exist_ok=True)
_cfg.PIPELINES_DIR.mkdir(parents=True, exist_ok=True)

# Every skull module below reads these as plain names at import time, so
# the redirected Paths above must be set before this import - not after.
from skull.tools import skills as _sm  # noqa: E402
_sm.SKILLS_DIR = _cfg.SKILLS_DIR
_sm.INDEX_PATH = _cfg.SKILLS_DIR / "index.json"

from skull.storage import store as _mem  # noqa: E402
_mem.MEMORY_DIR = _cfg.MEMORY_DIR

from skull.tools import pipeline as _pl  # noqa: E402
_pl.PIPELINES_DIR = _cfg.PIPELINES_DIR

from skull.core.session import Session  # noqa: E402

# write_file/run_command block on a real interactive y/n via input() - fine
# for the real REPL, but this harness runs non-interactively (no stdin to
# read), so ask_permission would always hit EOF and auto-deny. Scenario 4
# needs to verify a write actually happens once out of plan mode, so this
# harness auto-approves on its behalf - this is test-harness scaffolding
# for driving the approval gate deterministically, not a change to the gate
# itself (a real interactive session still prompts normally).
from skull.tools import files as _files  # noqa: E402
_files.ask_permission = lambda *a, **k: True

# Same reasoning as ask_permission above: request_skill_env blocks on a
# real getpass() prompt, which would hit EOF non-interactively. Scenario 6
# needs to actually set a value to verify get_env() picks it up, so this
# harness supplies one on the skill's behalf via a module-level variable
# the test sets before triggering the turn - never a real secret.
from skull.tools import skill_env as _skenv  # noqa: E402
_skenv.SKILLS_ENV_PATH = _TEMP_ROOT / "skills.env"
_SCENARIO_6_FAKE_SECRET_VALUE = ["not-set"]
_skenv.getpass.getpass = lambda prompt: _SCENARIO_6_FAKE_SECRET_VALUE[0]


class ScenarioFailure(AssertionError):
    pass


def _check(condition: bool, message: str):
    if not condition:
        raise ScenarioFailure(message)


def _new_session() -> Session:
    return Session()


def _tool_calls(session: Session) -> list:
    """Every tool name the model actually called across all turns so far,
    in order."""
    names = []
    for m in session.messages:
        if m.get("role") == "assistant" and m.get("tool_calls"):
            names.extend(tc["function"]["name"] for tc in m["tool_calls"])
    return names


def _last_assistant_text(session: Session) -> str:
    for m in reversed(session.messages):
        if m.get("role") == "assistant" and m.get("content"):
            return m["content"]
    return ""


# ---------------------------------------------------------------------------
# Scenario 1: document + skill chain
# ---------------------------------------------------------------------------

def scenario_1_document_and_skill_chain(workdir: Path):
    import docx

    from skull.tools import skills as sm

    sm.create_skill(
        "miles_to_km_s1", "convert miles to kilometers",
        {"type": "object", "properties": {"miles": {"type": "number"}}, "required": ["miles"]},
        "def run(**kwargs):\n    m = kwargs['miles']\n    return {'miles': m, 'km': round(m * 1.609344, 3)}\n",
    )
    sm.create_skill(
        "c_to_f_s1", "convert celsius to fahrenheit",
        {"type": "object", "properties": {"celsius": {"type": "number"}}, "required": ["celsius"]},
        "def run(**kwargs):\n    c = kwargs['celsius']\n    return {'celsius': c, 'fahrenheit': round(c * 9 / 5 + 32, 2)}\n",
    )

    doc = docx.Document()
    doc.add_heading("Site Survey", level=1)
    doc.add_paragraph("Distance from base camp: 15 miles.")
    doc.add_paragraph("Recorded temperature: 22 celsius.")
    doc_path = workdir / "survey.docx"
    doc.save(str(doc_path))

    session = _new_session()
    session.handle_turn(
        f"Read {doc_path} and convert the distance to kilometers and the temperature to "
        "fahrenheit, using your existing skills for the conversions. State both converted "
        "numbers clearly."
    )

    calls = _tool_calls(session)
    _check("read_file" in calls, f"expected read_file to be called, got: {calls}")
    _check(
        "miles_to_km_s1" in calls or "c_to_f_s1" in calls,
        f"expected at least one conversion skill to be called, got: {calls}",
    )

    text = _last_assistant_text(session)
    _check("24.14" in text or "24.1" in text, f"expected ~24.14 km in the reply, got: {text!r}")
    _check("71.6" in text, f"expected 71.6°F in the reply, got: {text!r}")


# ---------------------------------------------------------------------------
# Scenario 2: skill lifecycle (create, overwrite, rollback)
# ---------------------------------------------------------------------------

def scenario_2_skill_lifecycle(workdir: Path):
    from skull.tools import skills as sm

    session = _new_session()
    session.handle_turn(
        "Create a skill called scenario_square_s2 that takes a number n and returns n squared "
        "as {'n': n, 'squared': ...}. Then call it with n=6."
    )
    result_v1 = sm.run_skill("scenario_square_s2", {"n": 6})
    _check(
        result_v1 == {"result": {"n": 6, "squared": 36}},
        f"expected squared=36 after creation, got: {result_v1}",
    )

    session.handle_turn(
        "Now overwrite the scenario_square_s2 skill so it returns n cubed instead "
        "(as {'n': n, 'cubed': ...}). Then call it again with n=6 to confirm."
    )
    result_v2 = sm.run_skill("scenario_square_s2", {"n": 6})
    _check(
        result_v2 == {"result": {"n": 6, "cubed": 216}},
        f"expected cubed=216 after overwrite (this is the exact stale-bytecode bug shape "
        f"if it fails), got: {result_v2}",
    )

    session.handle_turn("Roll back scenario_square_s2 to its previous version.")
    result_v3 = sm.run_skill("scenario_square_s2", {"n": 6})
    _check(
        result_v3 == {"result": {"n": 6, "squared": 36}},
        f"expected squared=36 restored after rollback, got: {result_v3}",
    )


# ---------------------------------------------------------------------------
# Scenario 3: memory contradiction + fuzzy forget
# ---------------------------------------------------------------------------

def scenario_3_memory_contradiction_and_fuzzy_forget(workdir: Path):
    from skull.storage import store as mem

    session = _new_session()
    session.handle_turn("Remember that I like my code reviews to be extremely detailed and thorough.")
    session.handle_turn(
        "Actually, forget that - from now on I want code review feedback to be as brief and "
        "minimal as possible, just the critical issues."
    )

    persona_facts = [e["text"] for e in mem.persona().all()]
    detailed_still_present = any("detailed" in f.lower() or "thorough" in f.lower() for f in persona_facts)
    brief_present = any("brief" in f.lower() or "minimal" in f.lower() for f in persona_facts)
    _check(
        brief_present,
        f"expected the new brief-feedback preference to be stored, got facts: {persona_facts}",
    )
    _check(
        not detailed_still_present,
        f"expected the old detailed-feedback preference to be superseded (not left active "
        f"alongside the contradicting one), got facts: {persona_facts}",
    )

    session.handle_turn("Also remember that my favorite programming language is Rust.")
    session.handle_turn("Please forget the fact about my favorite programming language.")
    persona_facts_after = [e["text"] for e in mem.persona().all()]
    rust_still_present = any("rust" in f.lower() for f in persona_facts_after)
    _check(
        not rust_still_present,
        f"expected the Rust fact to be forgotten even without an exact quote (this is the "
        f"exact forget()-silent-failure-on-paraphrase bug shape if it fails), got: "
        f"{persona_facts_after}",
    )


# ---------------------------------------------------------------------------
# Scenario 4: plan-mode boundary
# ---------------------------------------------------------------------------

def scenario_4_plan_mode_boundary(workdir: Path):
    from skull.tools import skills as sm

    sm.create_skill(
        "double_s4", "doubles a number",
        {"type": "object", "properties": {"n": {"type": "integer"}}, "required": ["n"]},
        "def run(**kwargs):\n    return {'doubled': kwargs['n'] * 2}\n",
    )
    _check(
        sm.get_skill("double_s4")["side_effects"] == "read_only",
        "expected double_s4 to be classified read_only - if this fails, the classifier "
        "itself regressed, not just plan mode",
    )

    session = _new_session()
    session.plan_mode = True

    session.handle_turn("Use the double_s4 skill to double the number 21.")
    calls = _tool_calls(session)
    _check(
        "double_s4" in calls,
        f"expected the read-only skill to be callable in plan mode, got calls: {calls}",
    )
    text = _last_assistant_text(session)
    _check("42" in text, f"expected 42 in the reply, got: {text!r}")

    target_path = str(workdir / "should_not_exist.txt")
    session.handle_turn(
        f"Now write a file at {target_path} containing the word 'done'."
    )
    _check(
        not Path(target_path).exists(),
        "expected the write to NOT happen in plan mode - this is the exact "
        "wholesale-skill-exclusion/plan-mode-boundary bug shape if it fails",
    )
    write_calls = _tool_calls(session)
    _check(
        "write_file" not in write_calls[len(calls):],
        f"expected write_file to be absent from the model's tool list in plan mode, "
        f"but it was called: {write_calls}",
    )

    session.plan_mode = False
    session.handle_turn(f"OK, now actually write that file at {target_path} containing 'done'.")
    _check(
        Path(target_path).exists() and Path(target_path).read_text().strip() == "done",
        "expected the write to actually happen after switching to /auto",
    )


# ---------------------------------------------------------------------------
# Scenario 5: long/heavy turn
# ---------------------------------------------------------------------------

def scenario_5_long_heavy_turn(workdir: Path):
    from skull.core import compaction as comp

    big_path = workdir / "big_notes.txt"
    lines = [f"Line {i}: the quick brown fox jumps over the lazy dog, iteration {i}." for i in range(400)]
    big_path.write_text("\n".join(lines))

    session = _new_session()
    try:
        ok = session.handle_turn(
            f"Read {big_path} and tell me: how many lines does it have, and what is the "
            "text of line 200? Keep your answer to two short sentences."
        )
    except Exception as e:
        raise ScenarioFailure(f"turn raised an exception instead of handling gracefully: {e}")

    _check(ok, "expected handle_turn to report success (no 400/context-window failure)")

    for m in session.messages:
        _check(m.get("role") is not None, f"found a structurally invalid message: {m}")
    for i, m in enumerate(session.messages):
        if m.get("role") == "assistant" and m.get("tool_calls"):
            ids_needed = {tc["id"] for tc in m["tool_calls"]}
            found_ids = set()
            j = i + 1
            while j < len(session.messages) and session.messages[j].get("role") == "tool":
                found_ids.add(session.messages[j]["tool_call_id"])
                j += 1
            _check(
                ids_needed <= found_ids,
                f"tool_calls at index {i} missing a tool response - invalid conversation shape",
            )

    text = _last_assistant_text(session)
    _check(len(text) > 0, "expected a non-empty final reply")


# ---------------------------------------------------------------------------
# Scenario 6: skill credential never passes through the model
# ---------------------------------------------------------------------------

def scenario_6_skill_credential_stays_out_of_the_model(workdir: Path):
    """Modeled directly on a real transcript: a user asked the model to send
    an email, and the model asked the user to paste an SMTP password
    straight into the chat conversation - landing it in plain text in
    conversation history and long-term memory. Verifies the model instead
    declares the credential via required_env, calls request_skill_env
    (never seeing the value), and that no message anywhere in the
    conversation contains the fake secret value used here."""
    from skull.tools import skill_env as scenv
    from skull.tools import skills as sm

    fake_secret = "sk-not-a-real-secret-9f8e7d6c"
    _SCENARIO_6_FAKE_SECRET_VALUE[0] = fake_secret

    session = _new_session()
    session.handle_turn(
        "Create a skill called notify_s6 that sends a notification using an API key. "
        "It should take a 'message' parameter, but the API key must NOT be a plain "
        "parameter the user types into chat - use required_env for it (call the env var "
        "FAKE_NOTIFY_API_KEY), reading it via skull.tools.skill_env.get_env at call time. "
        "The skill can just return {'sent': True, 'had_key': key is not None} without "
        "actually calling any real API. After creating it, set up the credential."
    )

    entry = sm.get_skill("notify_s6")
    _check(entry is not None, "expected notify_s6 to have been created")
    _check(
        "FAKE_NOTIFY_API_KEY" in (entry.get("required_env") or []),
        f"expected FAKE_NOTIFY_API_KEY in required_env, got: {entry.get('required_env')}",
    )
    _check(
        "api_key" not in (entry.get("parameters") or {}).get("properties", {})
        and not any(
            "key" in str(v).lower() and "type" in v
            for v in (entry.get("parameters") or {}).get("properties", {}).values()
        ),
        "the credential must not have been declared as a plain parameter",
    )

    calls = _tool_calls(session)
    _check(
        "request_skill_env" in calls,
        f"expected request_skill_env to have been called to set up the credential, got: {calls}",
    )
    _check(
        scenv.get_env("FAKE_NOTIFY_API_KEY") == fake_secret,
        "expected the fake secret to have actually been stored via get_env",
    )

    # The critical safety property: the secret value must never appear
    # anywhere in the conversation the model can see.
    import json

    for m in session.messages:
        serialized = json.dumps(m)
        _check(
            fake_secret not in serialized,
            f"the secret value leaked into a conversation message: {m}",
        )

    # Confirm the skill's own code actually reads it correctly at runtime.
    result = sm.run_skill("notify_s6", {"message": "hello"})
    _check(
        result.get("result", {}).get("had_key") is True,
        f"expected the skill to see the credential via get_env() at call time, got: {result}",
    )


SCENARIOS = [
    ("Document + skill chain", scenario_1_document_and_skill_chain),
    ("Skill lifecycle (create/overwrite/rollback)", scenario_2_skill_lifecycle),
    ("Memory contradiction + fuzzy forget", scenario_3_memory_contradiction_and_fuzzy_forget),
    ("Plan-mode boundary", scenario_4_plan_mode_boundary),
    ("Long/heavy turn", scenario_5_long_heavy_turn),
    ("Skill credential stays out of the model", scenario_6_skill_credential_stays_out_of_the_model),
]


def main():
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--only", type=int, help="Run only scenario N (1-6)")
    args = parser.parse_args()

    if not _cfg.LLM_KEY:
        print("LLM_KEY is not set - these scenarios need a real live endpoint. Aborting.", file=sys.stderr)
        sys.exit(1)

    to_run = SCENARIOS if args.only is None else [SCENARIOS[args.only - 1]]

    results = []
    for i, (label, fn) in enumerate(to_run, start=1 if args.only is None else args.only):
        scenario_workdir = _TEMP_ROOT / f"scenario_{i}"
        scenario_workdir.mkdir(exist_ok=True)
        print(f"\n[{i}/{len(SCENARIOS)}] {label} ...")
        try:
            fn(scenario_workdir)
            print(f"  PASS")
            results.append((i, label, True, None))
        except ScenarioFailure as e:
            print(f"  FAIL: {e}")
            results.append((i, label, False, str(e)))
        except Exception as e:
            print(f"  ERROR: {type(e).__name__}: {e}")
            traceback.print_exc()
            results.append((i, label, False, f"{type(e).__name__}: {e}"))
        finally:
            shutil.rmtree(scenario_workdir, ignore_errors=True)

    print("\n" + "=" * 60)
    passed = sum(1 for *_, ok, _ in results if ok)
    for i, label, ok, err in results:
        status = "PASS" if ok else "FAIL"
        print(f"  [{status}] {i}. {label}")
    print(f"\n{passed}/{len(results)} scenarios passed")

    shutil.rmtree(_TEMP_ROOT, ignore_errors=True)
    sys.exit(0 if passed == len(results) else 1)


if __name__ == "__main__":
    main()
